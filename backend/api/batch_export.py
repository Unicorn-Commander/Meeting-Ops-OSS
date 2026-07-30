from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional, Dict, Any, Literal
from pydantic import BaseModel, Field
from datetime import datetime, timezone
from io import BytesIO
import base64
import binascii
import html
import logging
import json
import os
import re
import uuid
from pathlib import Path
from sqlalchemy.orm.attributes import flag_modified

from database.database import get_db
from auth.dependencies import get_current_organization, get_current_user
from auth.organization import ActiveOrganization
from auth.models import Organization, User
from auth.tier import gate_feature_for_caller
from database.models import RecordingSession

router = APIRouter(prefix="/api/export", tags=["Export"])
logger = logging.getLogger(__name__)


def _report_logo_path() -> Optional[str]:
    """Resolve the optional Meeting-Ops report mark without a network fetch."""
    candidates = [
        os.getenv("MEETING_OPS_REPORT_LOGO", "").strip(),
        "/app/assets/meeting-ops-mark.png",
        str(
            Path(__file__).resolve().parents[2]
            / "frontend"
            / "public"
            / "brand"
            / "meeting-ops-mark.png"
        ),
    ]
    return next((path for path in candidates if path and Path(path).is_file()), None)


BrandMode = Literal["default", "meeting_ops", "workspace", "unbranded"]
_BRAND_MODES = {"meeting_ops", "workspace", "unbranded"}
_HEX_COLOR_RE = re.compile(r"^#[0-9A-Fa-f]{6}$")
_LOGO_DATA_RE = re.compile(
    r"^data:image/(?P<kind>png|jpeg);base64,(?P<payload>[A-Za-z0-9+/=]+)$"
)
_MAX_REPORT_LOGO_BYTES = 512 * 1024


class ReportBrandingView(BaseModel):
    display_name: str
    accent_color: str
    default_mode: Literal["meeting_ops", "workspace", "unbranded"]
    has_logo: bool
    logo_data_uri: Optional[str] = None


class ReportBrandingUpdate(BaseModel):
    display_name: Optional[str] = Field(default=None, max_length=100)
    accent_color: Optional[str] = Field(default=None, max_length=7)
    default_mode: Optional[
        Literal["meeting_ops", "workspace", "unbranded"]
    ] = None
    logo_data_uri: Optional[str] = Field(default=None, max_length=720_000)
    clear_logo: bool = False


def _organization_report_settings(
    organization: Optional[Organization],
) -> dict[str, Any]:
    if organization is None:
        return {}
    settings = (
        organization.settings
        if isinstance(getattr(organization, "settings", None), dict)
        else {}
    )
    report = settings.get("report_branding")
    return dict(report) if isinstance(report, dict) else {}


def _validated_logo_bytes(data_uri: Optional[str]) -> Optional[bytes]:
    if not data_uri:
        return None
    match = _LOGO_DATA_RE.fullmatch(data_uri.strip())
    if not match:
        raise ValueError("Logo must be a PNG or JPEG data URL.")
    try:
        raw = base64.b64decode(match.group("payload"), validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ValueError("Logo data is not valid base64.") from exc
    if len(raw) > _MAX_REPORT_LOGO_BYTES:
        raise ValueError("Logo must be 512 KB or smaller.")
    kind = match.group("kind")
    if kind == "png" and not raw.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError("Logo content is not a valid PNG.")
    if kind == "jpeg" and not raw.startswith(b"\xff\xd8\xff"):
        raise ValueError("Logo content is not a valid JPEG.")
    try:
        from reportlab.lib.utils import ImageReader

        width, height = ImageReader(BytesIO(raw)).getSize()
    except Exception as exc:  # noqa: BLE001
        raise ValueError("Logo image could not be decoded.") from exc
    if width < 1 or height < 1 or width > 4096 or height > 4096:
        raise ValueError("Logo dimensions must be between 1 and 4096 pixels.")
    return raw


def _serialize_report_branding(
    organization: Organization,
) -> ReportBrandingView:
    report = _organization_report_settings(organization)
    accent = str(report.get("accent_color") or "#7C3AED")
    if not _HEX_COLOR_RE.fullmatch(accent):
        accent = "#7C3AED"
    mode = str(report.get("default_mode") or "meeting_ops")
    if mode not in _BRAND_MODES:
        mode = "meeting_ops"
    logo_data_uri = report.get("logo_data_uri")
    try:
        has_logo = _validated_logo_bytes(logo_data_uri) is not None
    except ValueError:
        logo_data_uri = None
        has_logo = False
    return ReportBrandingView(
        display_name=(
            str(report.get("display_name") or organization.name).strip()
            or organization.name
        ),
        accent_color=accent.upper(),
        default_mode=mode,  # type: ignore[arg-type]
        has_logo=has_logo,
        logo_data_uri=logo_data_uri if has_logo else None,
    )


def _require_report_brand_admin(
    active_org: ActiveOrganization,
    current_user: User,
) -> None:
    if current_user.is_superuser:
        return
    if active_org.role_name not in {"owner", "admin", "manager"}:
        raise HTTPException(
            status_code=403,
            detail="Workspace report branding requires an admin or manager.",
        )


@router.get("/branding", response_model=ReportBrandingView)
async def get_report_branding(
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
) -> ReportBrandingView:
    """Return safe workspace report-brand settings for previews/exports."""
    return _serialize_report_branding(active_org.organization)


@router.put("/branding", response_model=ReportBrandingView)
async def update_report_branding(
    payload: ReportBrandingUpdate,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
    db: Session = Depends(get_db),
) -> ReportBrandingView:
    """Persist a small, database-backed white-label lockup per workspace."""
    _require_report_brand_admin(active_org, current_user)
    organization = active_org.organization
    settings = dict(organization.settings or {})
    report = _organization_report_settings(organization)

    if payload.display_name is not None:
        report["display_name"] = (
            payload.display_name.strip() or organization.name
        )
    if payload.accent_color is not None:
        accent = payload.accent_color.strip()
        if not _HEX_COLOR_RE.fullmatch(accent):
            raise HTTPException(
                status_code=422,
                detail="Accent color must be a six-digit hex color such as #7C3AED.",
            )
        report["accent_color"] = accent.upper()
    if payload.default_mode is not None:
        report["default_mode"] = payload.default_mode
    if payload.clear_logo:
        report.pop("logo_data_uri", None)
    elif payload.logo_data_uri is not None:
        try:
            _validated_logo_bytes(payload.logo_data_uri)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        report["logo_data_uri"] = payload.logo_data_uri.strip()

    settings["report_branding"] = report
    organization.settings = settings
    flag_modified(organization, "settings")
    db.commit()
    db.refresh(organization)
    return _serialize_report_branding(organization)


def _report_organization(session: RecordingSession) -> Optional[Organization]:
    """Load the session workspace only when the ORM row is still bound."""
    try:
        from sqlalchemy.orm import object_session

        db = object_session(session)
        if db is None or not getattr(session, "organization_id", None):
            return None
        return (
            db.query(Organization)
            .filter(Organization.id == session.organization_id)
            .first()
        )
    except Exception:  # noqa: BLE001 - detached/test rows use safe defaults
        logger.debug("Could not resolve report organization", exc_info=True)
        return None


def _resolve_report_brand(
    session: RecordingSession,
    options: "ExportOptions",
) -> dict[str, Any]:
    organization = _report_organization(session)
    configured = (
        _serialize_report_branding(organization)
        if organization is not None
        else None
    )
    mode = options.brandMode
    if mode == "default":
        mode = configured.default_mode if configured else "meeting_ops"

    if mode == "unbranded":
        return {
            "mode": "unbranded",
            "name": "",
            "tagline": "",
            "accent_color": "#374151",
            "logo_path": None,
            "logo_bytes": None,
            "footer": "Meeting report",
        }

    if mode == "workspace":
        report = _organization_report_settings(organization)
        logo_data_uri = report.get("logo_data_uri")
        try:
            logo_bytes = _validated_logo_bytes(logo_data_uri)
        except ValueError:
            logo_bytes = None
        name = (
            configured.display_name
            if configured
            else (
                getattr(organization, "name", None)
                or "Meeting Intelligence"
            )
        )
        return {
            "mode": "workspace",
            "name": name,
            "tagline": "MEETING INTELLIGENCE REPORT",
            "accent_color": (
                configured.accent_color if configured else "#374151"
            ),
            "logo_path": None,
            "logo_bytes": logo_bytes,
            "footer": f"{name}  •  Meeting intelligence report",
        }

    return {
        "mode": "meeting_ops",
        "name": "MEETING-OPS",
        "tagline": "MEETING INTELLIGENCE REPORT",
        "accent_color": "#7C3AED",
        "logo_path": _report_logo_path(),
        "logo_bytes": None,
        "footer": "MEETING-OPS  •  Meeting intelligence, ready to share",
    }


def _brand_logo_source(brand: dict[str, Any]):
    if brand.get("logo_bytes"):
        return BytesIO(brand["logo_bytes"])
    return brand.get("logo_path")


# Export models
class ExportOptions(BaseModel):
    includeTimestamps: bool = True
    includeSpeakers: bool = True
    includeInsights: bool = False
    # ``includeTranscript`` gates whether ``export_to_markdown`` appends the
    # full diarized transcript after the summary sections. Pydantic v2
    # silently drops kwargs the model doesn't declare, so leaving this off
    # made every ``ExportOptions(includeTranscript=...)`` call quietly lose
    # the value and then crash with AttributeError at access time (the
    # silent-fail mode behind the Markdown export button).
    includeTranscript: bool = False
    # "default" resolves through the active workspace's saved choice.
    # White-label and unbranded modes never include Meeting-Ops vendor copy.
    brandMode: BrandMode = "default"
    mergeFiles: bool = False
    emailTo: Optional[str] = None
    scheduleTime: Optional[str] = None

class BatchExportRequest(BaseModel):
    sessionIds: List[str]
    format: str  # txt, pdf, docx, csv, json, srt
    options: ExportOptions

class ExportJob(BaseModel):
    id: str
    sessionIds: List[str]
    format: str
    status: str  # pending, processing, completed, failed
    progress: int
    createdAt: datetime
    completedAt: Optional[datetime] = None
    fileSize: Optional[int] = None
    downloadUrl: Optional[str] = None
    error: Optional[str] = None
    options: Dict[str, Any]
    # Email delivery outcome (only set when options.emailTo was requested).
    # Honest reporting: emailSent is True ONLY after a 2xx from Postmark;
    # any failure sets emailSent=False + emailError with the reason.
    emailSent: Optional[bool] = None
    emailError: Optional[str] = None

class ExportTemplate(BaseModel):
    id: str
    name: str
    format: str
    options: Dict[str, Any]
    description: str


def _get_session_for_export(
    db: Session,
    organization_id: int,
    session_identifier: str,
) -> Optional[RecordingSession]:
    session = db.query(RecordingSession).filter(
        RecordingSession.organization_id == organization_id,
        RecordingSession.session_id == session_identifier,
    ).first()
    if session:
        return session

    try:
        int_id = int(session_identifier)
    except (TypeError, ValueError):
        return None

    return db.query(RecordingSession).filter(
        RecordingSession.organization_id == organization_id,
        RecordingSession.id == int_id,
    ).first()


# ---------------------------------------------------------------------------
# Email delivery for batch exports — REAL send via Postmark, reusing the
# attachment-capable helper that already powers the per-meeting Email action
# (api.session_emails._postmark_send). Imported lazily because session_emails
# itself lazily imports renderers from this module (avoid a cycle).
# ---------------------------------------------------------------------------

# Postmark rejects messages over ~10MB total (attachment content counts
# base64-encoded, +~33%). Cap raw attachment bytes well below that.
_EMAIL_ATTACHMENT_MAX_BYTES = 7 * 1024 * 1024

_ATTACHMENT_CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".srt": "text/plain",
    ".json": "application/json",
    ".zip": "application/zip",
}


def _email_configured() -> bool:
    """True when the env carries everything _postmark_send needs."""
    from auth.email import _postmark_token

    return bool(_postmark_token())


def _send_export_email(job: "ExportJob", to_email: str, artifact_path: Path) -> None:
    """Email the finished export artifact to ``to_email`` and record the
    outcome on the job. Never raises; never reports success it didn't get."""
    from api.session_emails import _postmark_send  # lazy: avoid import cycle

    try:
        size = artifact_path.stat().st_size
        if size > _EMAIL_ATTACHMENT_MAX_BYTES:
            job.emailSent = False
            job.emailError = (
                f"Export is {size // (1024 * 1024)}MB — too large to email "
                f"(limit {_EMAIL_ATTACHMENT_MAX_BYTES // (1024 * 1024)}MB). "
                "Use the download link instead."
            )
            logger.warning("batch export %s: %s", job.id, job.emailError)
            return

        import base64

        content_type = _ATTACHMENT_CONTENT_TYPES.get(
            artifact_path.suffix.lower(), "application/octet-stream"
        )
        attachment = {
            "Name": artifact_path.name,
            "Content": base64.b64encode(artifact_path.read_bytes()).decode("ascii"),
            "ContentType": content_type,
        }

        n = len(job.sessionIds)
        subject = f"Your Meeting-Ops export ({n} meeting{'s' if n != 1 else ''}, {job.format})"
        html_body = (
            '<div style="font-family: Arial, sans-serif; color: #374151;">'
            "<p>Your Meeting-Ops batch export is ready and attached.</p>"
            f"<p><strong>{n}</strong> meeting{'s' if n != 1 else ''} &middot; "
            f"format <strong>{job.format}</strong> &middot; file "
            f"<strong>{artifact_path.name}</strong></p>"
            '<hr style="border: none; border-top: 1px solid #e5e7eb;" />'
            '<p style="color: #9ca3af; font-size: 12px;">Sent by Meeting-Ops.</p>'
            "</div>"
        )

        result = _postmark_send(
            to_email=to_email,
            to_name=None,
            subject=subject,
            html_body=html_body,
            attachments=[attachment],
        )
        if result.get("ok"):
            job.emailSent = True
            logger.info(
                "batch export %s emailed to %s (message_id=%s)",
                job.id, to_email, result.get("message_id"),
            )
        else:
            job.emailSent = False
            job.emailError = f"Email delivery failed: {result.get('error') or 'unknown error'}"
            logger.warning("batch export %s: %s", job.id, job.emailError)
    except Exception as exc:  # noqa: BLE001 — email must never sink the job
        job.emailSent = False
        job.emailError = f"Email delivery failed: {exc}"
        logger.warning("batch export %s: email send raised: %s", job.id, exc)


# In-memory storage for export jobs (should use database in production)
export_jobs: Dict[str, ExportJob] = {}
export_job_scope: Dict[str, Dict[str, int]] = {}
export_templates: List[ExportTemplate] = [
    ExportTemplate(
        id="1",
        name="Meeting Summary",
        format="pdf",
        options={"includeTimestamps": False, "includeSpeakers": True, "includeInsights": True},
        description="Professional meeting summary with key insights"
    ),
    ExportTemplate(
        id="2",
        name="Full Transcript",
        format="docx",
        options={"includeTimestamps": True, "includeSpeakers": True, "includeInsights": False},
        description="Complete transcript with timestamps and speakers"
    ),
    ExportTemplate(
        id="3",
        name="Subtitles",
        format="srt",
        options={"includeTimestamps": True, "includeSpeakers": False, "includeInsights": False},
        description="SRT subtitle format for video editing"
    )
]
export_template_scope: Dict[str, Dict[str, int]] = {}
BUILTIN_TEMPLATE_IDS = {template.id for template in export_templates}

@router.post("/batch", response_model=ExportJob)
async def create_batch_export(
    request: BatchExportRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
    db: Session = Depends(get_db)
):
    """Create a batch export job for multiple sessions"""
    # v3.18.3: defense-in-depth tier gate. Free users have no server-side
    # canonical content to export (browser-only tier); explicit gate
    # closes the surface in case server-side fallbacks ever land.
    gate_feature_for_caller(current_user, "canonical_reprocess", active_org)
    # Email delivery is real (Postmark, attachment included) when the env is
    # configured. When it is NOT configured we refuse up front with 501 —
    # never accept an emailTo job we know we cannot deliver.
    if request.options.emailTo:
        if not _email_configured():
            raise HTTPException(
                status_code=501,
                detail="Emailing batch exports is not available on this server "
                       "(email delivery is not configured). Download the file "
                       "instead, or use the per-meeting Email action.",
            )
        if "@" not in request.options.emailTo:
            raise HTTPException(
                status_code=422,
                detail="emailTo is not a valid email address.",
            )
    try:
        for session_id in request.sessionIds:
            if not _get_session_for_export(db, active_org.organization.id, session_id):
                raise HTTPException(
                    status_code=404,
                    detail="One or more sessions were not found in the active organization",
                )

        # Create export job
        job_id = str(uuid.uuid4())
        job = ExportJob(
            id=job_id,
            sessionIds=request.sessionIds,
            format=request.format,
            status="pending",
            progress=0,
            createdAt=datetime.now(timezone.utc),
            options=request.options.model_dump()
        )
        
        # Store job
        export_jobs[job_id] = job
        export_job_scope[job_id] = {
            "user_id": current_user.id,
            "organization_id": active_org.organization.id,
        }
        
        # Start background export task
        background_tasks.add_task(
            process_export_job,
            job_id,
            request.sessionIds,
            request.format,
            request.options,
            active_org.organization.id,
            db
        )
        
        logger.info(f"Created batch export job {job_id} for {len(request.sessionIds)} sessions")
        return job

    except HTTPException:
        # Don't let the generic handler below re-wrap deliberate 4xx (e.g.
        # the 404 for a session outside the active org) as a 500.
        raise
    except Exception as e:
        logger.error(f"Failed to create batch export: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/jobs", response_model=List[ExportJob])
async def get_export_jobs(
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Get all export jobs for the current user"""
    return [
        job
        for job_id, job in export_jobs.items()
        if export_job_scope.get(job_id, {}).get("user_id") == current_user.id
        and export_job_scope.get(job_id, {}).get("organization_id") == active_org.organization.id
    ]

@router.get("/jobs/{job_id}", response_model=ExportJob)
async def get_export_job(
    job_id: str,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Get status of a specific export job"""
    job = export_jobs.get(job_id)
    scope = export_job_scope.get(job_id, {})
    if not job or scope.get("user_id") != current_user.id or scope.get("organization_id") != active_org.organization.id:
        raise HTTPException(status_code=404, detail="Export job not found")
    return job

@router.delete("/jobs/{job_id}")
async def cancel_export_job(
    job_id: str,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Cancel an export job"""
    job = export_jobs.get(job_id)
    scope = export_job_scope.get(job_id, {})
    if not job or scope.get("user_id") != current_user.id or scope.get("organization_id") != active_org.organization.id:
        raise HTTPException(status_code=404, detail="Export job not found")
    
    if job.status in ["completed", "failed"]:
        raise HTTPException(status_code=400, detail="Cannot cancel completed or failed job")
    
    job.status = "cancelled"
    job.error = "Job cancelled by user"
    
    return {"message": "Export job cancelled"}

@router.get("/templates", response_model=List[ExportTemplate])
async def get_export_templates(
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Get available export templates"""
    return [
        template
        for template in export_templates
        if template.id in BUILTIN_TEMPLATE_IDS
        or export_template_scope.get(template.id, {}).get("organization_id") == active_org.organization.id
    ]

@router.post("/templates", response_model=ExportTemplate)
async def create_export_template(
    template: ExportTemplate,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Create a new export template"""
    # v3.18.3: tier gate templates so free users can't seed configs that
    # would only get exercised by the gated batch endpoint.
    gate_feature_for_caller(current_user, "canonical_reprocess", active_org)
    template.id = str(uuid.uuid4())
    export_templates.append(template)
    export_template_scope[template.id] = {
        "user_id": current_user.id,
        "organization_id": active_org.organization.id,
    }
    return template

@router.delete("/templates/{template_id}")
async def delete_export_template(
    template_id: str,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Delete an export template"""
    global export_templates
    scope = export_template_scope.get(template_id, {})
    if template_id not in BUILTIN_TEMPLATE_IDS and (
        scope.get("organization_id") != active_org.organization.id
        or scope.get("user_id") != current_user.id
    ):
        raise HTTPException(status_code=404, detail="Template not found")
    export_templates = [t for t in export_templates if t.id != template_id]
    export_template_scope.pop(template_id, None)
    return {"message": "Template deleted"}

async def process_export_job(
    job_id: str,
    session_ids: List[str],
    format: str,
    options: ExportOptions,
    organization_id: int,
    db: Session
):
    """Background task to process export job"""
    job = export_jobs.get(job_id)
    if not job:
        return
    
    try:
        job.status = "processing"
        total_sessions = len(session_ids)
        
        # Create export directory
        export_dir = Path("/tmp/exports") / job_id
        export_dir.mkdir(parents=True, exist_ok=True)
        
        exported_files = []
        
        for i, session_id in enumerate(session_ids):
            # Update progress
            job.progress = int((i / total_sessions) * 100)
            
            # Get session from database
            session = _get_session_for_export(db, organization_id, session_id)
            
            if not session:
                logger.warning(f"Session {session_id} not found")
                continue
            
            # Export based on format
            if format == "txt":
                content = export_to_txt(session, options)
                filename = f"{session.name or session_id}.txt"
            elif format == "json":
                content = export_to_json(session, options)
                filename = f"{session.name or session_id}.json"
            elif format == "srt":
                content = export_to_srt(session, options)
                filename = f"{session.name or session_id}.srt"
            elif format == "pdf":
                pdf_bytes = export_to_pdf(session, options)
                filename = f"{session.name or session_id}.pdf"
                file_path = export_dir / filename
                with open(file_path, 'wb') as f:
                    f.write(pdf_bytes)
                exported_files.append(file_path)
                continue
            elif format == "docx":
                docx_bytes = export_to_docx(session, options)
                filename = f"{session.name or session_id}.docx"
                file_path = export_dir / filename
                with open(file_path, 'wb') as f:
                    f.write(docx_bytes)
                exported_files.append(file_path)
                continue
            else:
                content = f"# {session.name or session_id}\n\nExport format {format} not supported."
                filename = f"{session.name or session_id}.txt"

            # Save text-based file
            file_path = export_dir / filename
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            exported_files.append(file_path)
        
        # Create archive if multiple files or merge requested
        artifact_path: Optional[Path] = None
        if len(exported_files) > 1 or options.mergeFiles:
            import zipfile
            archive_path = export_dir / f"export_{job_id}.zip"
            with zipfile.ZipFile(archive_path, 'w') as zf:
                for file_path in exported_files:
                    zf.write(file_path, file_path.name)

            job.downloadUrl = f"/api/export/download/{job_id}/export_{job_id}.zip"
            job.fileSize = archive_path.stat().st_size
            artifact_path = archive_path
        elif exported_files:
            # Single file
            job.downloadUrl = f"/api/export/download/{job_id}/{exported_files[0].name}"
            job.fileSize = exported_files[0].stat().st_size
            artifact_path = exported_files[0]

        job.status = "completed"
        job.progress = 100
        job.completedAt = datetime.now(timezone.utc)

        # Real email delivery (Postmark + attachment) when requested.
        # _send_export_email records the outcome on the job (emailSent /
        # emailError) and never raises — an email failure must not flip a
        # successfully-generated export to "failed" (download still works).
        if options.emailTo:
            if artifact_path is not None:
                _send_export_email(job, options.emailTo, artifact_path)
            else:
                job.emailSent = False
                job.emailError = "Export produced no file to email."
                logger.warning(
                    "batch export %s: emailTo=%s requested but no artifact was produced",
                    job_id, options.emailTo,
                )

    except Exception as e:
        logger.error(f"Export job {job_id} failed: {e}")
        job.status = "failed"
        job.error = str(e)

def _get_summary_data(session: RecordingSession) -> Dict[str, Any]:
    """Extract summary data from session, handling both legacy and new formats."""
    summary_data = {
        "executive": "",
        "bullets": [],
        "actions": [],
        "decisions": [],
        "participants": (
            session.participants
            if isinstance(getattr(session, "participants", None), list)
            else []
        ),
        "title": session.title or session.name or "Untitled Meeting",
    }

    # Try final_summary first (preferred)
    final_loaded = False
    final = session.final_summary
    if final:
        if isinstance(final, str):
            try:
                final = json.loads(final)
            except (json.JSONDecodeError, TypeError):
                final = {}
        if isinstance(final, dict):
            final_loaded = True
            summary_data["executive"] = final.get("executive", "")
            summary_data["bullets"] = final.get("bullets", [])
            summary_data["actions"] = final.get("actions", [])
            summary_data["decisions"] = final.get("decisions", [])
            if final.get("title"):
                summary_data["title"] = final["title"]
            # Continue below so first-class ActionItem rows can replace the
            # stale action snapshot captured at summarization time.

    # Fallback to legacy summary field
    legacy = session.summary
    if legacy and not final_loaded:
        if isinstance(legacy, str):
            try:
                legacy = json.loads(legacy)
            except (json.JSONDecodeError, TypeError):
                legacy = {}
        if isinstance(legacy, dict):
            # Could be nested under 'analysis' key
            source = legacy.get("analysis", legacy)
            summary_data["executive"] = source.get("executive", "")
            summary_data["bullets"] = source.get("bullets", [])
            summary_data["actions"] = source.get("actions", [])
            summary_data["decisions"] = source.get("decisions", [])
            if source.get("title"):
                summary_data["title"] = source["title"]

    # The editable ActionItem table is canonical after summarization.  When
    # the ORM row is bound, render current owners/due dates/statuses in every
    # report instead of the frozen final_summary.actions list.
    try:
        from sqlalchemy.orm import object_session
        from database.models import ActionItem

        db = object_session(session)
        if db is not None:
            current_actions = (
                db.query(ActionItem)
                .filter(
                    ActionItem.session_id == session.id,
                    ActionItem.organization_id == session.organization_id,
                )
                .order_by(ActionItem.sort_order.asc(), ActionItem.id.asc())
                .all()
            )
            if current_actions:
                summary_data["actions"] = [
                    {
                        "action": item.text,
                        "owner": item.owner or "",
                        "status": item.status or "todo",
                        "due_date": (
                            item.due_date.strftime("%b %d, %Y")
                            if item.due_date
                            else ""
                        ),
                        "priority": (item.raw_payload or {}).get("priority", ""),
                    }
                    for item in current_actions
                ]
    except Exception:  # noqa: BLE001 - a report should still use summary data
        logger.exception("Could not load canonical action items for report")

    return summary_data


def _get_transcript_segments(session: RecordingSession) -> list:
    """Extract transcript segments from session data."""
    # Try diarized transcript first
    if session.transcript_diarized and isinstance(session.transcript_diarized, dict):
        # v3.44: hydrate speaker names live from the current profiles so every
        # export format (PDF/DOCX/TXT/SRT/JSON/markdown) reflects a rename
        # without rewriting the stored transcript.
        from services.speaker_service import hydrate_diarized_for_session
        hydrated = hydrate_diarized_for_session(session)
        segments = (hydrated or {}).get("segments", []) if isinstance(hydrated, dict) else []
        if segments:
            return segments

    # Try transcript JSON field
    if session.transcript:
        try:
            data = json.loads(session.transcript) if isinstance(session.transcript, str) else session.transcript
            if isinstance(data, dict) and "segments" in data:
                return data["segments"]
        except (json.JSONDecodeError, TypeError):
            pass

    return []


def _get_plain_transcript(session: RecordingSession) -> str:
    """Get plain text transcript from session."""
    if session.transcript_simple:
        return session.transcript_simple
    if session.transcript:
        try:
            data = json.loads(session.transcript) if isinstance(session.transcript, str) else session.transcript
            if isinstance(data, dict) and "text" in data:
                return data["text"]
            if isinstance(data, str):
                return data
        except (json.JSONDecodeError, TypeError):
            return str(session.transcript)
    return ""


def export_to_pdf(session: RecordingSession, options: ExportOptions) -> bytes:
    """Export session to a professional PDF document using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image, SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

    brand = _resolve_report_brand(session, options)
    buffer = BytesIO()
    page_width, page_height = letter
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=1.0 * inch,
        bottomMargin=0.8 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )

    # Colors
    purple = HexColor(brand["accent_color"])
    dark_gray = HexColor("#374151")
    medium_gray = HexColor("#6B7280")
    light_gray = HexColor("#F3F4F6")
    yellow_bg = HexColor("#FEF3C7")
    green_text = HexColor("#065F46")

    # Custom styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "MeetingTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        textColor=dark_gray,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=purple,
        spaceBefore=16,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        "BodyText2",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        textColor=dark_gray,
        leading=14,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "BulletText",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        textColor=dark_gray,
        leading=14,
        leftIndent=16,
        spaceAfter=4,
        bulletIndent=0,
    ))
    styles.add(ParagraphStyle(
        "MetaText",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        textColor=medium_gray,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "TranscriptSpeaker",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=9,
        textColor=purple,
        spaceAfter=1,
    ))
    styles.add(ParagraphStyle(
        "TranscriptText",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        textColor=dark_gray,
        leading=12,
        spaceAfter=8,
        leftIndent=8,
    ))
    styles.add(ParagraphStyle(
        "BrandMark",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        textColor=HexColor("#FFFFFF"),
    ))
    styles.add(ParagraphStyle(
        "BrandLockup",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=13,
        textColor=dark_gray,
    ))
    styles.add(ParagraphStyle(
        "TableText",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=dark_gray,
    ))
    styles.add(ParagraphStyle(
        "TableHeaderText",
        parent=styles["TableText"],
        fontName="Helvetica-Bold",
        textColor=HexColor("#FFFFFF"),
    ))

    story = []
    summary_data = _get_summary_data(session)

    # --- Branded report lockup ---
    if brand["mode"] != "unbranded":
        safe_brand_name = html.escape(str(brand["name"]))
        safe_tagline = html.escape(str(brand["tagline"]))
        brand_text = Paragraph(
            f"{safe_brand_name}<br/><font size=\"7\" "
            f'color="{brand["accent_color"]}">{safe_tagline}</font>',
            styles["BrandLockup"],
        )
        logo_source = _brand_logo_source(brand)
        fallback_mark = brand["mode"] == "meeting_ops" and not logo_source
        if logo_source:
            brand_mark = Image(logo_source)
            brand_mark._restrictSize(0.72 * inch, 0.54 * inch)
            lockup_rows = [[brand_mark, brand_text]]
            lockup_widths = [0.82 * inch, 3.2 * inch]
        elif fallback_mark:
            lockup_rows = [[Paragraph("M", styles["BrandMark"]), brand_text]]
            lockup_widths = [0.42 * inch, 3.2 * inch]
        else:
            lockup_rows = [[brand_text]]
            lockup_widths = [3.8 * inch]

        brand_lockup = Table(lockup_rows, colWidths=lockup_widths)
        brand_lockup_style = [
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (0, 0), 0 if logo_source else 4),
            ("RIGHTPADDING", (0, 0), (0, 0), 4),
            ("TOPPADDING", (0, 0), (0, 0), 0 if logo_source else 4),
            ("BOTTOMPADDING", (0, 0), (0, 0), 0 if logo_source else 4),
        ]
        if fallback_mark:
            brand_lockup_style.extend([
                ("BACKGROUND", (0, 0), (0, 0), purple),
                ("LEFTPADDING", (1, 0), (1, 0), 9),
            ])
        elif logo_source:
            brand_lockup_style.append(
                ("LEFTPADDING", (1, 0), (1, 0), 9)
            )
        brand_lockup.setStyle(TableStyle(brand_lockup_style))
        story.append(brand_lockup)
    story.append(HRFlowable(
        width="100%",
        thickness=2 if brand["mode"] != "unbranded" else 0.75,
        color=purple,
        spaceAfter=12,
    ))
    story.append(Paragraph(
        summary_data["title"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
        styles["MeetingTitle"]
    ))

    # Meeting metadata uses the user-corrected meeting date rather than the
    # ingestion timestamp whenever it is available.
    meeting_day = (
        session.meeting_date
        or (session.started_at.date() if session.started_at else None)
        or (session.created_at.date() if session.created_at else None)
    )
    date_str = meeting_day.strftime("%B %d, %Y") if meeting_day else "N/A"
    if session.meeting_time:
        date_str += f" at {session.meeting_time.strftime('%I:%M %p')}"
    duration_secs = session.duration or 0
    duration_str = f"{int(duration_secs // 60)}m {int(duration_secs % 60)}s" if duration_secs else "N/A"
    participants = summary_data["participants"]
    meta_table = Table(
        [[
            Paragraph(f"<b>Date</b><br/>{date_str}", styles["MetaText"]),
            Paragraph(f"<b>Duration</b><br/>{duration_str}", styles["MetaText"]),
            Paragraph(
                f"<b>Participants</b><br/>{len(participants) if participants else 'Not listed'}",
                styles["MetaText"],
            ),
            Paragraph(
                "<b>Transcript</b><br/>"
                + ("Included" if options.includeTranscript else "Not included"),
                styles["MetaText"],
            ),
        ]],
        colWidths=[2.25 * inch, 1.3 * inch, 1.45 * inch, 1.5 * inch],
    )
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), light_gray),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor("#E5E7EB")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, HexColor("#E5E7EB")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 12))

    if participants:
        participant_names = []
        for participant in participants:
            if isinstance(participant, dict):
                participant_names.append(
                    participant.get("name")
                    or participant.get("display_name")
                    or participant.get("email")
                    or "Participant"
                )
            else:
                participant_names.append(str(participant))
        story.append(Paragraph(
            "<b>Participants:</b> "
            + ", ".join(
                name.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                for name in participant_names
            ),
            styles["MetaText"],
        ))

    has_summary = bool(
        summary_data["executive"] or summary_data["bullets"]
        or summary_data["actions"] or summary_data["decisions"]
    )

    # --- Executive Summary ---
    if summary_data["executive"]:
        story.append(Paragraph("Executive Summary", styles["SectionHeading"]))
        safe_text = summary_data["executive"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe_text, styles["BodyText2"]))
        story.append(Spacer(1, 8))

    # --- Key Discussion Points ---
    if summary_data["bullets"]:
        story.append(Paragraph("Key Discussion Points", styles["SectionHeading"]))
        for bullet in summary_data["bullets"]:
            safe = str(bullet).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(f"\u2022  {safe}", styles["BulletText"]))
        story.append(Spacer(1, 8))

    # --- Action Items ---
    if summary_data["actions"]:
        story.append(Paragraph("Action Items", styles["SectionHeading"]))
        action_rows = [[
            Paragraph("Action", styles["TableHeaderText"]),
            Paragraph("Owner", styles["TableHeaderText"]),
            Paragraph("Status", styles["TableHeaderText"]),
            Paragraph("Due", styles["TableHeaderText"]),
        ]]
        for action in summary_data["actions"]:
            if isinstance(action, dict):
                action_text = action.get("action", str(action))
                owner = action.get("owner", "")
                status = action.get("status", "todo")
                due = action.get("due_date", "")
            else:
                action_text = str(action)
                owner = ""
                status = "todo"
                due = ""
            safe_action = str(action_text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            safe_owner = str(owner).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            action_rows.append([
                Paragraph(safe_action, styles["TableText"]),
                Paragraph(safe_owner or "—", styles["TableText"]),
                Paragraph(str(status).replace("_", " ").title(), styles["TableText"]),
                Paragraph(str(due) or "—", styles["TableText"]),
            ])
        action_table = Table(
            action_rows,
            colWidths=[3.35 * inch, 1.15 * inch, 0.85 * inch, 1.15 * inch],
            repeatRows=1,
        )
        action_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), purple),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
            ("GRID", (0, 0), (-1, -1), 0.35, HexColor("#D1D5DB")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), light_gray]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(action_table)
        story.append(Spacer(1, 8))

    # --- Important Decisions ---
    if summary_data["decisions"]:
        story.append(Paragraph("Important Decisions", styles["SectionHeading"]))
        for decision in summary_data["decisions"]:
            safe = str(decision).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(f"\u2713  {safe}", styles["BulletText"]))
        story.append(Spacer(1, 8))

    # --- Optional transcript appendix ---
    segments = _get_transcript_segments(session) if options.includeTranscript else []
    plain_transcript = _get_plain_transcript(session) if options.includeTranscript else ""

    if options.includeTranscript and (segments or plain_transcript):
        if has_summary:
            story.append(PageBreak())
        story.append(Paragraph("Full Transcript", styles["SectionHeading"]))

        if segments:
            for seg in segments:
                speaker = seg.get("speaker", "Speaker")
                text = seg.get("text", "")
                start = seg.get("start", 0)

                label_parts = []
                if options.includeSpeakers:
                    safe_speaker = speaker.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    label_parts.append(safe_speaker)
                if options.includeTimestamps:
                    label_parts.append(f"[{format_timestamp(start)}]")

                if label_parts:
                    story.append(Paragraph(" ".join(label_parts), styles["TranscriptSpeaker"]))

                safe_text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_text, styles["TranscriptText"]))
        elif plain_transcript:
            # Wrap long plain text into paragraphs
            for para in plain_transcript.split("\n"):
                if para.strip():
                    safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe, styles["TranscriptText"]))
    elif not has_summary:
        message = (
            "No summary data is available for this session."
            if not options.includeTranscript
            else "No transcript or summary data is available for this session."
        )
        story.append(Paragraph(message, styles["BodyText2"]))

    # --- Footer callback ---
    def add_footer(canvas, doc):
        canvas.saveState()
        # Purple accent line
        canvas.setStrokeColor(purple)
        canvas.setLineWidth(0.5)
        canvas.line(0.75 * inch, 0.6 * inch, page_width - 0.75 * inch, 0.6 * inch)
        # Footer text
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(medium_gray)
        canvas.drawString(
            0.75 * inch,
            0.45 * inch,
            str(brand["footer"]),
        )
        canvas.drawRightString(page_width - 0.75 * inch, 0.45 * inch, f"Page {doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    return buffer.getvalue()


def export_to_docx(session: RecordingSession, options: ExportOptions) -> bytes:
    """Export session to a professional DOCX document using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    brand = _resolve_report_brand(session, options)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    # --- Style configuration ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    accent_bytes = bytes.fromhex(brand["accent_color"].lstrip("#"))
    purple_rgb = RGBColor(*accent_bytes)
    dark_gray_rgb = RGBColor(0x37, 0x41, 0x51)
    medium_gray_rgb = RGBColor(0x6B, 0x72, 0x80)
    for heading_name in ("Title", "Heading 1", "Heading 2"):
        heading_style = doc.styles[heading_name]
        heading_style.font.name = "Calibri"
        heading_style.font.color.rgb = (
            dark_gray_rgb if heading_name in {"Title", "Heading 1"} else purple_rgb
        )

    summary_data = _get_summary_data(session)

    # --- Brand header ---
    if brand["mode"] != "unbranded":
        brand_para = doc.add_paragraph()
        brand_para.paragraph_format.space_after = Pt(1)
        logo_source = _brand_logo_source(brand)
        if logo_source:
            logo_run = brand_para.add_run()
            logo_run.add_picture(logo_source, width=Inches(0.5))
            brand_para.add_run("  ")
        brand_run = brand_para.add_run(str(brand["name"]))
        brand_run.font.size = Pt(12)
        brand_run.font.bold = True
        brand_run.font.color.rgb = purple_rgb
        report_run = brand_para.add_run(f"\n{brand['tagline']}")
        report_run.font.size = Pt(7)
        report_run.font.bold = True
        report_run.font.color.rgb = medium_gray_rgb

    # Purple divider line (thin horizontal rule via a border paragraph)
    divider = doc.add_paragraph()
    divider_fmt = divider.paragraph_format
    divider_fmt.space_after = Pt(6)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), brand["accent_color"].lstrip("#"))
    pBdr.append(bottom)
    divider.paragraph_format.element.get_or_add_pPr().append(pBdr)

    # --- Title ---
    title_para = doc.add_heading(summary_data["title"], level=1)
    for run in title_para.runs:
        run.font.color.rgb = dark_gray_rgb

    # --- Metadata ---
    meeting_day = (
        session.meeting_date
        or (session.started_at.date() if session.started_at else None)
        or (session.created_at.date() if session.created_at else None)
    )
    date_str = meeting_day.strftime("%B %d, %Y") if meeting_day else "N/A"
    if session.meeting_time:
        date_str += f" at {session.meeting_time.strftime('%I:%M %p')}"
    duration_secs = session.duration or 0
    duration_str = f"{int(duration_secs // 60)}m {int(duration_secs % 60)}s" if duration_secs else "N/A"
    participants = summary_data["participants"]

    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_values = [
        ("DATE", date_str),
        ("DURATION", duration_str),
        ("PARTICIPANTS", str(len(participants)) if participants else "Not listed"),
        ("TRANSCRIPT", "Included" if options.includeTranscript else "Not included"),
    ]
    for cell, (label, value) in zip(meta_table.rows[0].cells, meta_values):
        cell.text = ""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F3F0FF")
        cell._tc.get_or_add_tcPr().append(shading)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(label + "\n")
        label_run.font.size = Pt(7)
        label_run.font.bold = True
        label_run.font.color.rgb = purple_rgb
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(8)
        value_run.font.color.rgb = dark_gray_rgb

    if participants:
        participant_names = []
        for participant in participants:
            if isinstance(participant, dict):
                participant_names.append(
                    participant.get("name")
                    or participant.get("display_name")
                    or participant.get("email")
                    or "Participant"
                )
            else:
                participant_names.append(str(participant))
        participant_para = doc.add_paragraph()
        participant_para.paragraph_format.space_before = Pt(5)
        participant_para.paragraph_format.space_after = Pt(2)
        participant_label = participant_para.add_run("Participants: ")
        participant_label.font.bold = True
        participant_label.font.size = Pt(8)
        participant_label.font.color.rgb = medium_gray_rgb
        participant_run = participant_para.add_run(", ".join(participant_names))
        participant_run.font.size = Pt(8)
        participant_run.font.color.rgb = medium_gray_rgb

    has_summary = bool(
        summary_data["executive"] or summary_data["bullets"]
        or summary_data["actions"] or summary_data["decisions"]
    )

    # --- Executive Summary ---
    if summary_data["executive"]:
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(summary_data["executive"])

    # --- Key Discussion Points ---
    if summary_data["bullets"]:
        doc.add_heading("Key Discussion Points", level=2)
        for bullet in summary_data["bullets"]:
            doc.add_paragraph(str(bullet), style="List Bullet")

    # --- Action Items (Table) ---
    if summary_data["actions"]:
        doc.add_heading("Action Items", level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Task", "Owner", "Status", "Due"]
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        for action in summary_data["actions"]:
            if isinstance(action, dict):
                task = action.get("action", str(action))
                owner = action.get("owner", "")
                status = action.get("status", "Pending")
                due_date = action.get("due_date", "")
            else:
                task = str(action)
                owner = ""
                status = "Pending"
                due_date = ""

            row_cells = table.add_row().cells
            row_cells[0].text = task
            row_cells[1].text = owner or "—"
            row_cells[2].text = str(status).replace("_", " ").title()
            row_cells[3].text = due_date or "—"

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # --- Important Decisions ---
    if summary_data["decisions"]:
        doc.add_heading("Important Decisions", level=2)
        for decision in summary_data["decisions"]:
            doc.add_paragraph(str(decision), style="List Bullet")

    # --- Optional transcript appendix ---
    segments = _get_transcript_segments(session) if options.includeTranscript else []
    plain_transcript = _get_plain_transcript(session) if options.includeTranscript else ""

    if options.includeTranscript and (segments or plain_transcript):
        if has_summary:
            doc.add_page_break()
        doc.add_heading("Full Transcript", level=2)

        if segments:
            for seg in segments:
                speaker = seg.get("speaker", "Speaker")
                text = seg.get("text", "")
                start = seg.get("start", 0)

                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)

                label_parts = []
                if options.includeSpeakers:
                    label_parts.append(speaker)
                if options.includeTimestamps:
                    label_parts.append(f"[{format_timestamp(start)}]")

                if label_parts:
                    label_run = p.add_run(" ".join(label_parts) + ": ")
                    label_run.font.bold = True
                    label_run.font.size = Pt(9)
                    label_run.font.color.rgb = purple_rgb

                text_run = p.add_run(text)
                text_run.font.size = Pt(9)
        elif plain_transcript:
            for para_text in plain_transcript.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.space_after = Pt(4)
                    for run in p.runs:
                        run.font.size = Pt(9)
    elif not has_summary:
        doc.add_paragraph(
            "No summary data is available for this session."
            if not options.includeTranscript
            else "No transcript or summary data is available for this session."
        )

    # --- Footer ---
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(str(brand["footer"]))
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = medium_gray_rgb

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_to_txt(session: RecordingSession, options: ExportOptions) -> str:
    """Export session to plain text format"""
    lines = []
    lines.append(f"Meeting: {session.name or 'Untitled'}")
    lines.append(f"Date: {session.created_at}")
    lines.append(f"Duration: {session.duration or 0} seconds")
    lines.append("-" * 50)
    lines.append("")
    
    # Add transcript if available
    if hasattr(session, 'transcript') and session.transcript:
        transcript = json.loads(session.transcript) if isinstance(session.transcript, str) else session.transcript
        
        for segment in transcript.get('segments', []):
            line = ""
            
            if options.includeTimestamps:
                start_time = segment.get('start', 0)
                line += f"[{format_timestamp(start_time)}] "
            
            if options.includeSpeakers and 'speaker' in segment:
                line += f"{segment['speaker']}: "
            
            line += segment.get('text', '')
            lines.append(line)
    else:
        lines.append("No transcript available for this session.")
    
    if options.includeInsights:
        lines.append("")
        lines.append("-" * 50)
        lines.append("AI INSIGHTS")
        lines.append("-" * 50)
        lines.append("• Key topics discussed")
        lines.append("• Action items identified")
        lines.append("• Decisions made")
    
    return "\n".join(lines)

def export_to_markdown(session: RecordingSession, options: ExportOptions) -> str:
    """Render the session summary as a GitHub-flavored markdown document.

    Pulls structured fields from final_summary first, then falls back to
    parsing session.summary (JSON-stringified). Includes transcript only
    when options.includeTranscript is set.
    """
    brand = _resolve_report_brand(session, options)
    title = session.title or session.name or "Untitled Meeting"
    duration_min = round((session.duration or 0) / 60, 1)
    created = session.created_at.strftime("%Y-%m-%d %H:%M") if session.created_at else "(unknown date)"

    # Pull the summary payload — final_summary is canonical (JSON column).
    fs = session.final_summary if isinstance(session.final_summary, dict) else {}
    if not fs and session.summary:
        try:
            fs = json.loads(session.summary) if isinstance(session.summary, str) else session.summary
        except Exception:
            fs = {}

    lines: list[str] = []
    if brand["mode"] != "unbranded":
        lines.append(f"> **{brand['name']}** · {brand['tagline'].title()}")
        lines.append("")
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"**Date:** {created}  ")
    lines.append(f"**Duration:** {duration_min} minutes")
    if session.description:
        lines.append(f"**Notes:** {session.description}")
    lines.append("")

    executive = (fs.get("executive") or "").strip()
    if executive:
        lines.append("## Executive Summary")
        lines.append("")
        lines.append(executive)
        lines.append("")

    bullets = fs.get("bullets") or []
    if isinstance(bullets, list) and bullets:
        lines.append("## Key Discussion Points")
        lines.append("")
        for b in bullets:
            if isinstance(b, str) and b.strip():
                lines.append(f"- {b.strip()}")
        lines.append("")

    actions = fs.get("actions") or fs.get("action_items") or []
    if isinstance(actions, list) and actions:
        lines.append("## Action Items")
        lines.append("")
        for a in actions:
            if isinstance(a, str):
                lines.append(f"- [ ] {a}")
            elif isinstance(a, dict):
                text = a.get("action") or a.get("text") or ""
                owner = a.get("owner") or a.get("assignee")
                priority = a.get("priority")
                suffix = []
                if owner:
                    suffix.append(f"Owner: {owner}")
                if priority:
                    suffix.append(f"Priority: {priority}")
                tail = f" ({', '.join(suffix)})" if suffix else ""
                lines.append(f"- [ ] {text}{tail}")
        lines.append("")

    decisions = fs.get("decisions") or fs.get("key_decisions") or []
    if isinstance(decisions, list) and decisions:
        lines.append("## Decisions")
        lines.append("")
        for d in decisions:
            if isinstance(d, str):
                lines.append(f"- {d}")
            elif isinstance(d, dict):
                lines.append(f"- {d.get('text') or d.get('decision') or json.dumps(d)}")
        lines.append("")

    questions = fs.get("questions") or fs.get("open_questions") or fs.get("follow_ups") or []
    if isinstance(questions, list) and questions:
        lines.append("## Open Questions")
        lines.append("")
        for q in questions:
            if isinstance(q, str):
                lines.append(f"- {q}")
        lines.append("")

    minutes = (fs.get("minutes") or "").strip()
    if minutes:
        lines.append("## Meeting Minutes")
        lines.append("")
        lines.append(minutes)
        lines.append("")

    # Optional transcript
    if options.includeTranscript:
        lines.append("## Transcript")
        lines.append("")
        # v3.44: render speaker names live from the current profiles (rename
        # shows instantly; no stored-transcript rewrite).
        from services.speaker_service import hydrate_diarized_for_session
        hydrated = hydrate_diarized_for_session(session)
        diarized = hydrated if isinstance(hydrated, dict) else None
        segs = diarized.get("segments") if diarized else None
        if segs:
            for seg in segs:
                t = seg.get("text") or ""
                if not t.strip():
                    continue
                spk = seg.get("speaker") or ""
                ts = ""
                if options.includeTimestamps and "start" in seg:
                    ts = f"[{format_timestamp(seg['start'])}] "
                speaker = f"**{spk}:** " if spk and options.includeSpeakers else ""
                lines.append(f"{ts}{speaker}{t.strip()}")
                lines.append("")
        elif session.transcript_simple:
            lines.append(session.transcript_simple)
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def export_to_json(session: RecordingSession, options: ExportOptions) -> str:
    """Export session to JSON format"""
    data = {
        "session_id": session.id,
        "name": session.name,
        "created_at": session.created_at.isoformat() if session.created_at else None,
        "duration": session.duration,
        "transcript": []
    }
    
    if hasattr(session, 'transcript') and session.transcript:
        transcript = json.loads(session.transcript) if isinstance(session.transcript, str) else session.transcript
        
        for segment in transcript.get('segments', []):
            segment_data = {"text": segment.get('text', '')}
            
            if options.includeTimestamps:
                segment_data["start"] = segment.get('start', 0)
                segment_data["end"] = segment.get('end', 0)
            
            if options.includeSpeakers and 'speaker' in segment:
                segment_data["speaker"] = segment['speaker']
            
            data["transcript"].append(segment_data)
    
    if options.includeInsights:
        data["insights"] = {
            "topics": [],
            "action_items": [],
            "decisions": []
        }
    
    return json.dumps(data, indent=2)

def export_to_srt(session: RecordingSession, options: ExportOptions) -> str:
    """Export session to SRT subtitle format"""
    lines = []
    counter = 1
    
    if hasattr(session, 'transcript') and session.transcript:
        transcript = json.loads(session.transcript) if isinstance(session.transcript, str) else session.transcript
        
        for segment in transcript.get('segments', []):
            lines.append(str(counter))
            
            start_time = format_srt_timestamp(segment.get('start', 0))
            end_time = format_srt_timestamp(segment.get('end', segment.get('start', 0) + 5))
            lines.append(f"{start_time} --> {end_time}")
            
            text = segment.get('text', '')
            if options.includeSpeakers and 'speaker' in segment:
                text = f"{segment['speaker']}: {text}"
            
            lines.append(text)
            lines.append("")  # Empty line between subtitles
            counter += 1
    
    return "\n".join(lines)

def format_timestamp(seconds: float) -> str:
    """Format seconds to HH:MM:SS"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"

def format_srt_timestamp(seconds: float) -> str:
    """Format seconds to SRT timestamp format (HH:MM:SS,mmm)"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:06.3f}".replace('.', ',')

@router.get("/download/{job_id}/{filename}")
async def download_export(
    job_id: str,
    filename: str,
    current_user: User = Depends(get_current_user),
    active_org: ActiveOrganization = Depends(get_current_organization),
):
    """Download exported file"""
    from fastapi.responses import FileResponse
    
    scope = export_job_scope.get(job_id, {})
    if scope.get("user_id") != current_user.id or scope.get("organization_id") != active_org.organization.id:
        raise HTTPException(status_code=404, detail="Export job not found")

    file_path = Path("/tmp/exports") / job_id / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type='application/octet-stream'
    )
